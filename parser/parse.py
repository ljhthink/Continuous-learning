#!/usr/bin/env python3
"""
P4 Phase 4b — 文件解析管道（Tauri sidecar）

用法:
    python parse.py <file_path> [--domain <domain>]

支持格式:
    PDF  → pymupdf (fitz) 提取文本 + 基本结构
    DOCX → python-docx 提取段落 + 表格
    XLSX → openpyxl 提取工作表 + 单元格
    MD   → 直接透传

输出:
    JSON 到 stdout，格式：
    {
        "success": true,
        "format": "pdf",
        "markdown": "...",
        "title": "从内容提取的标题",
        "metadata": { "pages": 10, "author": "..." }
    }
    或
    {
        "success": false,
        "error": "错误描述"
    }

退出码:
    0 = 成功
    1 = 解析失败
    2 = 不支持的格式
"""

import sys
import os
import json
import argparse
from pathlib import Path


def parse_pdf(file_path: str) -> dict:
    """用 pymupdf (fitz) 解析 PDF"""
    import fitz  # pymupdf

    doc = fitz.open(file_path)
    pages = []
    metadata = {
        "pages": len(doc),
        "author": doc.metadata.get("author", ""),
        "title": doc.metadata.get("title", ""),
    }

    for page_num, page in enumerate(doc, 1):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"## 第 {page_num} 页\n\n{text.strip()}")
        # 提取页面中的表格（简单文本表格）
        tables = page.find_tables()
        for table_idx, table in enumerate(tables):
            rows = table.extract()
            if rows:
                pages.append(f"\n### 表格 {table_idx + 1}\n")
                # markdown 表格
                header = rows[0]
                pages.append("| " + " | ".join(str(c or "") for c in header) + " |")
                pages.append("| " + " | ".join("---" for _ in header) + " |")
                for row in rows[1:]:
                    pages.append("| " + " | ".join(str(c or "") for c in row) + " |")

    doc.close()

    markdown = "\n\n".join(pages) if pages else "(空文档)"
    title = metadata.get("title") or Path(file_path).stem

    return {
        "success": True,
        "format": "pdf",
        "markdown": markdown,
        "title": title,
        "metadata": metadata,
    }


def parse_docx(file_path: str) -> dict:
    """用 python-docx 解析 DOCX"""
    from docx import Document

    doc = Document(file_path)
    lines = []
    paragraph_count = 0
    table_count = 0

    # 按文档顺序遍历段落和表格
    from docx.oxml.ns import qn

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            # 段落
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue
            paragraph_count += 1
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name.lower()
            if "heading 1" in style:
                lines.append(f"# {text}")
            elif "heading 2" in style:
                lines.append(f"## {text}")
            elif "heading 3" in style:
                lines.append(f"### {text}")
            elif "heading 4" in style:
                lines.append(f"#### {text}")
            elif "title" in style:
                lines.append(f"# {text}")
            else:
                lines.append(text)
        elif child.tag == qn("w:tbl"):
            # 表格
            table_count += 1
            # 找到对应的 Table 对象
            for t in doc.tables:
                if t._element is child:
                    rows = t.rows
                    if rows:
                        header = [cell.text.strip() for cell in rows[0].cells]
                        lines.append(f"\n| {' | '.join(header)} |")
                        lines.append(f"| {' | '.join('---' for _ in header)} |")
                        for row in rows[1:]:
                            cells = [cell.text.strip() for cell in row.cells]
                            lines.append(f"| {' | '.join(cells)} |")
                    break

    markdown = "\n\n".join(lines) if lines else "(空文档)"
    title = Path(file_path).stem

    # 尝试从 core properties 获取标题
    try:
        cp = doc.core_properties
        if cp.title:
            title = cp.title
    except Exception:
        pass

    return {
        "success": True,
        "format": "docx",
        "markdown": markdown,
        "title": title,
        "metadata": {
            "paragraphs": paragraph_count,
            "tables": table_count,
        },
    }


def parse_xlsx(file_path: str) -> dict:
    """用 openpyxl 解析 XLSX"""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    lines = []
    sheet_count = 0

    for sheet_name in wb.sheetnames:
        sheet_count += 1
        ws = wb[sheet_name]
        lines.append(f"# {sheet_name}\n")
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            row_count += 1
            # 跳过全空行
            if all(cell is None or str(cell).strip() == "" for cell in row):
                continue
            cells = [str(cell).strip() if cell is not None else "" for cell in row]
            if row_count == 1:
                lines.append(f"| {' | '.join(cells)} |")
                lines.append(f"| {' | '.join('---' for _ in cells)} |")
            else:
                lines.append(f"| {' | '.join(cells)} |")
        lines.append("")

    wb.close()
    markdown = "\n".join(lines) if lines else "(空工作簿)"
    title = Path(file_path).stem

    return {
        "success": True,
        "format": "xlsx",
        "markdown": markdown,
        "title": title,
        "metadata": {"sheets": sheet_count},
    }


def parse_markdown(file_path: str) -> dict:
    """直接透传 markdown 文件"""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    # 从第一行提取标题
    title = Path(file_path).stem
    for line in content.split("\n"):
        line = line.strip()
        if line.startswith("# "):
            title = line[2:].strip()
            break
        elif line.startswith("title:"):
            # frontmatter title
            title = line.split(":", 1)[1].strip().strip('"').strip("'")
            break

    return {
        "success": True,
        "format": "md",
        "markdown": content,
        "title": title,
        "metadata": {},
    }


# 格式 -> 解析函数映射
PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "xlsx": parse_xlsx,
    "md": parse_markdown,
}


def get_format(file_path: str) -> str:
    """从文件扩展名获取格式"""
    ext = Path(file_path).suffix.lower().lstrip(".")
    if ext == "txt":
        ext = "md"  # txt 当 md 处理
    return ext


def main():
    parser = argparse.ArgumentParser(description="P4 文件解析管道（Tauri sidecar）")
    parser.add_argument("file_path", help="要解析的文件路径")
    parser.add_argument("--domain", default="", help="领域分类（可选）")
    args = parser.parse_args()

    file_path = args.file_path

    # 验证文件存在
    if not os.path.exists(file_path):
        print(json.dumps({
            "success": False,
            "error": f"文件不存在: {file_path}",
        }, ensure_ascii=False))
        sys.exit(1)

    # 获取格式
    fmt = get_format(file_path)

    if fmt not in PARSERS:
        print(json.dumps({
            "success": False,
            "error": f"不支持的格式: .{fmt}（支持: pdf/docx/xlsx/md）",
        }, ensure_ascii=False))
        sys.exit(2)

    # 解析
    try:
        result = PARSERS[fmt](file_path)
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(0)
    except ImportError as e:
        print(json.dumps({
            "success": False,
            "error": f"依赖库缺失: {e}",
        }, ensure_ascii=False))
        sys.exit(1)
    except Exception as e:
        print(json.dumps({
            "success": False,
            "error": f"解析失败: {type(e).__name__}: {e}",
        }, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()
